"""
Resume Parser Module: Extracts text from PDF and DOCX files.
Uses pdfplumber (PDF), python-docx (DOCX), and Gemini 2.5 Flash (OCR fallback).
"""

import os
from typing import Optional
import pdfplumber
from docx import Document
import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))


def extract_text_from_pdf_with_gemini(file_path: str) -> Optional[str]:
    """Fallback OCR using Gemini 2.5 Flash for image-based PDFs."""
    try:
        print(f"  → Using Gemini 2.5 Flash for OCR...")
        model = genai.GenerativeModel("gemini-2.5-flash")
        uploaded_file = genai.upload_file(file_path)
        
        prompt = "Extract ALL text content from this resume document. Return ONLY the raw text."
        
        response = model.generate_content(
            [uploaded_file, prompt],
            generation_config=genai.GenerationConfig(temperature=0)
        )
        
        uploaded_file.delete()
        return response.text.strip() if response.text else None
        
    except Exception as e:
        print(f"  Error with Gemini OCR: {e}")
        return None


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """Extract text from PDF using pdfplumber, with Gemini OCR fallback."""
    try:
        text_content = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                if (page_text := page.extract_text()):
                    text_content.append(page_text)
        
        if text_content:
            return "\n".join(text_content)
        
        print(f"No text layer found, falling back to Gemini OCR...")
        return extract_text_from_pdf_with_gemini(file_path)
    
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return extract_text_from_pdf_with_gemini(file_path)


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """Extract text from DOCX files (paragraphs and tables)."""
    try:
        doc = Document(file_path)
        text_content = [p.text for p in doc.paragraphs if p.text.strip()]
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
            
        return "\n".join(text_content) if text_content else None
    
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return None


def parse_resume(file_path: str) -> Optional[str]:
    """Parse resume file (PDF/DOCX) and extract text."""
    if not os.path.exists(file_path):
        return None
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        print(f"Error: Unsupported file format: {file_ext}")
        return None


def parse_resumes(file_paths: list) -> dict:
    """Batch process list of resume files."""
    results = {}
    for file_path in file_paths:
        print(f"Parsing: {os.path.basename(file_path)}...")
        results[file_path] = parse_resume(file_path)
    return results


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        sys.exit("Usage: python parser.py <resume_file>")
    
    text = parse_resume(sys.argv[1])
    if text:
        print(f"\nEXTRACTED TEXT:\n{'='*50}\n{text[:2000]}")
